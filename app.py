from flask import Flask, render_template, request, send_file, flash, redirect
from docx import Document
from docx.shared import Cm
from docx.enum.section import WD_ORIENT
from io import BytesIO  # Importa BytesIO para manipulação de arquivos em memória

app = Flask(__name__)
app.secret_key = 'your_secret_key'  # Necessário para flash messages

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        # Verifica se 'image_count' está presente na requisição
        image_count = request.form.get('image_count')
        if not image_count:
            flash('Por favor, selecione o número de imagens.')
            return redirect(request.url)

        # Salva as imagens enviadas pelo usuário
        image_files = request.files.getlist('image_files')
        image_streams = []  # Lista para armazenar os arquivos de imagem em memória

        for image_file in image_files:
            if image_file and image_file.filename:
                # Armazena a imagem na memória usando BytesIO
                image_stream = BytesIO()
                image_file.save(image_stream)
                image_stream.seek(0)  # Move o ponteiro para o início do stream
                image_streams.append(image_stream)

        # Verifica se o número correto de imagens foi carregado
        if len(image_streams) != int(image_count):
            flash(f'Por favor, carregue exatamente {image_count} imagem(s).')
            return redirect(request.url)

        # Cria o documento Word
        doc = Document()

        # Configura layout
        if image_count == '1':
            # Layout retrato para uma imagem
            section = doc.sections[0]
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)

            # Define margens estreitas
            narrow_margin = Cm(1.27)
            section.top_margin = narrow_margin
            section.bottom_margin = narrow_margin
            section.left_margin = narrow_margin
            section.right_margin = narrow_margin

            # Insere a única imagem
            doc.add_paragraph().add_run().add_picture(image_streams[0], width=section.page_width - narrow_margin * 2, height=section.page_height - narrow_margin * 2)

        elif image_count == '2' and len(image_streams) == 2:
            # Layout paisagem para duas imagens
            section = doc.sections[0]
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Cm(29.7)
            section.page_height = Cm(21.0)

            # Define margens estreitas
            narrow_margin = Cm(1.27)
            section.top_margin = narrow_margin
            section.bottom_margin = narrow_margin
            section.left_margin = narrow_margin
            section.right_margin = narrow_margin

            # Adiciona uma tabela com uma linha e duas colunas para as imagens lado a lado
            height = Cm(17.52)
            width = Cm(13.45)

            table = doc.add_table(rows=1, cols=2)
            table.allow_autofit = False

            # Insere a primeira imagem
            cell_1 = table.cell(0, 0)
            cell_1.width = width
            cell_1.paragraphs[0].add_run().add_picture(image_streams[0], width=width, height=height)

            # Insere a segunda imagem
            cell_2 = table.cell(0, 1)
            cell_2.width = width
            cell_2.paragraphs[0].add_run().add_picture(image_streams[1], width=width, height=height)

        # Salva o documento em um objeto BytesIO para enviar como download
        doc_stream = BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)  # Move o ponteiro para o início do stream

        return send_file(doc_stream, as_attachment=True, download_name='Tarefa - [Formatada].docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    return render_template('upload_form.html')

if __name__ == '__main__':
    app.run(debug=True)
